import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pygments import lex
from pygments.lexers import get_lexer_by_name, CSharpLexer
from pygments.token import Token

import requests
from io import BytesIO
from PIL import Image

from site_photos import PAGES_IMAGES

try:
    from site_descriptions import PAGES_DESCRIPTION
except:
    PAGES_DESCRIPTION = {}

# =========================
# הגדרות
# =========================
project_path = r"C:\Users\aradl\source\repos\TheWorldOfHorses🐎"
output_file = "all_pages_documentation.docx"

pages = {}

# =========================
# פונקציות עזר
# =========================


def add_description(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(30, 30, 30)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_bg(paragraph, color="EDEDED"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def add_code(doc, text, lang="cs"):
    try:
        lexer = CSharpLexer() if lang == "cs" else get_lexer_by_name("html")
    except:
        lexer = get_lexer_by_name("text")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_bg(p, "F2F2F2")

    for token, value in lex(text, lexer):
        run = p.add_run(value)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

        color = RGBColor(40, 40, 40)

        if token in Token.Keyword:
            color = RGBColor(0, 0, 255)
        elif token in Token.String:
            color = RGBColor(163, 21, 21)
        elif token in Token.Comment:
            color = RGBColor(0, 128, 0)
        elif token in Token.Name:
            color = RGBColor(124, 91, 0)

        run.font.color.rgb = color


def safe_add_image(doc, img_bytes, i):
    try:
        img = Image.open(img_bytes)
        img = img.convert("RGB")

        fixed = BytesIO()
        img.save(fixed, format="PNG")
        fixed.seek(0)

        doc.add_picture(fixed, width=Inches(5))

    except Exception as e:
        print(f"⚠ Word נכשל בתמונה {i}: {repr(e)}")


# =========================
# סריקה
# =========================
if __name__ == "__main__":
    print("🔍 סורק פרויקט...")

    for root, _, files in os.walk(project_path):
        if any(skip in root for skip in [".git", ".vs", "bin", "obj"]):
            continue

        for file in files:
            if not (file.endswith(".aspx") or file.endswith(".aspx.cs")):
                continue

            path = os.path.join(root, file)

            page_name = (
                file.replace(".aspx.cs", "").replace(".aspx", "").replace(".cs", "")
            )

            if page_name not in pages:
                pages[page_name] = {"aspx": "", "cs": ""}

            try:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
            except:
                with open(path, "r", encoding="windows-1255") as f:
                    content = f.read()

            if file.endswith(".aspx") and not file.endswith(".aspx.cs"):
                pages[page_name]["aspx"] = content
            elif file.endswith(".aspx.cs"):
                pages[page_name]["cs"] = content

    # =========================
    # יצירת מסמך
    # =========================
    doc = Document()
    doc.add_heading("📘 TheWorldOfHorses - תיעוד פרויקט", 0)

    doc.add_heading("📍 ניווט מהיר", 1)
    for page in pages.keys():
        p = doc.add_paragraph(page)
        p.style = "List Bullet"

    doc.add_page_break()

    # =========================
    # כתיבה
    # =========================
    print("📄 כותב למסמך...")

    for page_name, data in pages.items():

        doc.add_heading(f"📄 {page_name}", 1)

        base = page_name

        desc = (
            PAGES_DESCRIPTION.get(page_name)
            or PAGES_DESCRIPTION.get(page_name + ".aspx")
            or PAGES_DESCRIPTION.get(base)
            or PAGES_DESCRIPTION.get(base + ".aspx")
            or "⚠ אין תיאור"
        )

        add_description(doc, "📝 " + desc)

        # =========================
        # תמונות
        # =========================
        images = PAGES_IMAGES.get(page_name) or PAGES_IMAGES.get(page_name + ".aspx")

        if images:
            doc.add_heading("🖼 תמונות", 2)

            for i, img in enumerate(images, start=1):
                try:
                    headers = {"User-Agent": "Mozilla/5.0"}
                    response = requests.get(img, headers=headers, timeout=25)
                    response.raise_for_status()

                    if not response.headers.get("Content-Type", "").startswith("image"):
                        raise Exception("לא תמונה אמיתית")

                    img_bytes = BytesIO(response.content)
                    safe_add_image(doc, img_bytes, i)

                except Exception as e:
                    print(f"⚠ תמונה נכשלה {i}: {repr(e)}")

        # =========================
        # ASPX
        # =========================
        if data["aspx"]:
            doc.add_heading("ASPX", 2)
            add_code(doc, data["aspx"], "html")

        # =========================
        # C#
        # =========================
        if data["cs"]:
            doc.add_heading("C#", 2)
            add_code(doc, data["cs"], "cs")

        doc.add_page_break()

    # =========================
    # שמירה
    # =========================
    doc.save(output_file)

    print("✅ סיום! נוצר קובץ:", output_file)
    os.system(f'start "" "{output_file}"')
